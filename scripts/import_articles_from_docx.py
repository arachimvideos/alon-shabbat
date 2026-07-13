from __future__ import annotations

import argparse
import html
import re
import sqlite3
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DB_PATH = ROOT / "backend" / "data" / "articles.db"


@dataclass
class Article:
    parasha: str
    title: str
    subtitle: str
    issue: str
    body: str
    tags: list[str]


def now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u200f", "")).strip()


def style_name(paragraph) -> str:
    return paragraph.style.name if paragraph.style else ""


def parasha_from_heading(value: str) -> str:
    name = clean_text(re.sub(r"\s*\(\d+\)\s*$", "", value))
    return canonical_parasha_name(name)


def canonical_parasha_name(name: str) -> str:
    name = re.sub(r"^פרשת\s+", "", name)
    name = re.sub(r"\s*[–-]\s*", " ", name)
    name = clean_text(name)
    aliases = {
        "תצווה": "תצוה",
        "קורח": "קרח",
        "חוקת": "חקת",
        "פנחס": "פינחס",
        "נצבים": "ניצבים",
        "נצבים וילך": "ניצבים וילך",
        "ויקהל פקודי": "ויקהל פקודי",
        "תזריע מצורע": "תזריע מצורע",
        "אחרי מות קדושים": "אחרי מות קדושים",
        "בהר בחוקותי": "בהר בחוקותי",
        "מטות מסעי": "מטות מסעי",
        "נשא חג השבועות": "נשא חג השבועות",
        "חג הפסח": "פסח",
        "בהר טז אייר": "בהר",
    }
    return aliases.get(name, name)


def issue_number_from_heading(value: str) -> str:
    match = re.search(r"\d+", value)
    return match.group(0) if match else clean_text(value)


def paragraph_to_html(style: str, text: str) -> str:
    escaped = html.escape(text)
    if style == "Heading 5":
        return f"<h2>{escaped}</h2>"
    return f"<p>{escaped}</p>"


def tags_from_text(value: str) -> list[str]:
    value = re.sub(r"^תגיות\s*:?", "", value).strip()
    if not value:
        return []
    parts = re.split(r"[,;،\n]+", value)
    tags: list[str] = []
    seen: set[str] = set()
    for part in parts:
        tag = clean_text(part)
        key = tag.casefold()
        if tag and key not in seen:
            seen.add(key)
            tags.append(tag)
    return tags


def parse_docx(path: Path) -> list[Article]:
    document = Document(str(path))
    starts = [
        index
        for index, paragraph in enumerate(document.paragraphs)
        if style_name(paragraph) == "Heading 1" and clean_text(paragraph.text)
    ]

    articles: list[Article] = []
    for position, start in enumerate(starts):
        end = starts[position + 1] if position + 1 < len(starts) else len(document.paragraphs)
        paragraphs = document.paragraphs[start:end]

        parasha = parasha_from_heading(paragraphs[0].text)
        title = ""
        subtitle = ""
        issue = ""
        body_parts: list[str] = []
        tags: list[str] = []

        for paragraph in paragraphs[1:]:
            text = clean_text(paragraph.text)
            if not text:
                continue

            style = style_name(paragraph)
            if style == "Heading 2" and not title:
                title = text
                continue
            if style == "Heading 3" and not subtitle:
                subtitle = text
                continue
            if style == "Heading 4" and not issue:
                issue = issue_number_from_heading(text)
                continue
            if style == "תגיות":
                tags.extend(tags_from_text(text))
                continue

            body_parts.append(paragraph_to_html(style, text))

        if not parasha or not title or not issue or not body_parts:
            continue

        body = "\n".join(body_parts)
        articles.append(Article(parasha=parasha, title=title, subtitle=subtitle, issue=issue, body=body, tags=tags))

    return articles


def ensure_schema(conn: sqlite3.Connection) -> None:
    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS parashot (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            created_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS tags (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            created_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS articles (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            subtitle TEXT,
            issue_number TEXT,
            parasha_id INTEGER NOT NULL REFERENCES parashot(id),
            author_name TEXT,
            publication_date TEXT,
            uploaded_at TEXT NOT NULL,
            body_text TEXT,
            extracted_text TEXT,
            status TEXT NOT NULL DEFAULT 'פורסם',
            original_filename TEXT,
            stored_filename TEXT,
            file_path TEXT,
            file_type TEXT
        );

        CREATE TABLE IF NOT EXISTS article_tags (
            article_id INTEGER NOT NULL REFERENCES articles(id) ON DELETE CASCADE,
            tag_id INTEGER NOT NULL REFERENCES tags(id),
            PRIMARY KEY (article_id, tag_id)
        );
        """
    )
    columns = {row[1] for row in conn.execute("PRAGMA table_info(articles)").fetchall()}
    migrations = {
        "subtitle": "ALTER TABLE articles ADD COLUMN subtitle TEXT",
        "issue_number": "ALTER TABLE articles ADD COLUMN issue_number TEXT",
    }
    for column, statement in migrations.items():
        if column not in columns:
            conn.execute(statement)


def get_or_create_parasha(conn: sqlite3.Connection, name: str) -> int:
    timestamp = now()
    conn.execute("INSERT OR IGNORE INTO parashot (name, created_at) VALUES (?, ?)", (name, timestamp))
    row = conn.execute("SELECT id FROM parashot WHERE name = ?", (name,)).fetchone()
    if row is None:
        raise RuntimeError(f"Could not create parasha: {name}")
    return int(row[0])


def get_or_create_tag(conn: sqlite3.Connection, name: str) -> int:
    timestamp = now()
    conn.execute("INSERT OR IGNORE INTO tags (name, created_at) VALUES (?, ?)", (name, timestamp))
    row = conn.execute("SELECT id FROM tags WHERE name = ?", (name,)).fetchone()
    if row is None:
        raise RuntimeError(f"Could not create tag: {name}")
    return int(row[0])


def existing_article_count(conn: sqlite3.Connection, article: Article, parasha_id: int) -> int:
    row = conn.execute(
        """
        SELECT COUNT(*)
        FROM articles
        WHERE title = ? AND parasha_id = ? AND COALESCE(body_text, '') = ?
        """,
        (article.title, parasha_id, article.body),
    ).fetchone()
    return int(row[0])


def import_articles(articles: list[Article]) -> tuple[int, int]:
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    inserted = 0
    skipped = 0
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("PRAGMA foreign_keys = ON")
        ensure_schema(conn)

        seen_this_import: dict[tuple[int, str, str], int] = {}
        for article in articles:
            parasha_id = get_or_create_parasha(conn, article.parasha)
            fingerprint = (parasha_id, article.title, article.body)
            seen_this_import[fingerprint] = seen_this_import.get(fingerprint, 0) + 1
            if existing_article_count(conn, article, parasha_id) >= seen_this_import[fingerprint]:
                skipped += 1
                continue

            conn.execute(
                """
                INSERT INTO articles (
                    title, subtitle, issue_number, parasha_id, author_name, publication_date, uploaded_at,
                    body_text, extracted_text, status, original_filename, stored_filename,
                    file_path, file_type
                )
                VALUES (?, ?, ?, ?, NULL, NULL, ?, ?, ?, 'פורסם', NULL, NULL, NULL, NULL)
                """,
                (article.title, article.subtitle or None, article.issue or None, parasha_id, now(), article.body, article.body),
            )
            article_id = int(conn.execute("SELECT last_insert_rowid()").fetchone()[0])
            for tag in article.tags:
                tag_id = get_or_create_tag(conn, tag)
                conn.execute(
                    "INSERT OR IGNORE INTO article_tags (article_id, tag_id) VALUES (?, ?)",
                    (article_id, tag_id),
                )
            inserted += 1

    return inserted, skipped


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    articles = parse_docx(args.docx)
    if args.dry_run:
        print(f"parsed={len(articles)}")
        for article in articles[:5]:
            print(f"{article.parasha} | {article.title} | {article.issue}")
        return

    inserted, skipped = import_articles(articles)
    print(f"parsed={len(articles)} inserted={inserted} skipped_duplicates={skipped}")


if __name__ == "__main__":
    main()
