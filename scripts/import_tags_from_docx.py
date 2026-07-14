from __future__ import annotations

import argparse
import os
import re
import sqlite3
from dataclasses import dataclass
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DB_PATH = ROOT / "backend" / "data" / "articles.db"

TAG_CATEGORIES = [
    "מוסר וערכים",
    "מוסר אנושי",
    "כח הדמיון",
    "פסיכולוגיה",
    "חינוך",
    "מידות",
    "מועדים",
    "מצוות",
    "אמונה",
]

TAG_FIXES = {
    "מתמדדים": "מתמודדים",
    "התפיסות": "התפייסות",
    "קינאה": "קנאה",
}

PARASHA_ALIASES = {
    "תצווה": "תצוה",
    "קורח": "קרח",
    "חוקת": "חקת",
    "פנחס": "פינחס",
    "נצבים": "ניצבים",
    "נצבים וילך": "ניצבים וילך",
    "ויקהל פקודי": "ויקהל פקודי",
    "תזריע מצורע": "תזריע מצורע",
    "אחרי מות קדושים": "אחרי מות קדושים",
    "בהר בחוקותי": "בהר בחוקותי",
    "מטות מסעי": "מטות מסעי",
    "נשא חג השבועות": "נשא חג השבועות",
    "חג הפסח": "פסח",
}


@dataclass
class TagRow:
    parasha: str
    title_text: str
    tags: list[str]


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def clean_env_value(name: str) -> str:
    value = (os.getenv(name) or "").strip()
    if (value.startswith('"') and value.endswith('"')) or (value.startswith("'") and value.endswith("'")):
        value = value[1:-1].strip()
    return value


def clean_text(value: str | None) -> str:
    return re.sub(r"\s+", " ", (value or "").replace("\u200f", "")).strip()


def normalize_for_match(value: str | None) -> str:
    value = clean_text(value).lower()
    value = re.sub(r"[\"״׳'`’‘.,:;?!()\[\]\-–—]", " ", value)
    return clean_text(value)


def parasha_from_cell(value: str) -> str:
    value = re.sub(r"\s*\(\d+\)\s*$", "", clean_text(value))
    value = re.sub(r"^פרשת\s+", "", value)
    value = re.sub(r"\s*[–-]\s*", " ", value)
    value = clean_text(value)
    return PARASHA_ALIASES.get(value, value)


def normalize_tag(value: str) -> str:
    value = clean_text(value.strip("[]() "))
    value = TAG_FIXES.get(value, value)
    return value


def expand_tag_part(value: str) -> list[str]:
    tag = normalize_tag(value)
    if not tag:
        return []
    for category in ["חינוך", "מידות", "מועדים", "מצוות", "פסיכולוגיה", "אמונה"]:
        prefix = f"{category} "
        if tag.startswith(prefix):
            return [category, normalize_tag(tag[len(prefix) :])]
    return [tag]


def split_tag_cell(value: str) -> list[str]:
    value = clean_text(value)
    if not value:
        return []

    category_pattern = "|".join(re.escape(category) for category in TAG_CATEGORIES)
    value = re.sub(rf"\s+(?=({category_pattern})(?:\b|$))", "\n", value)
    value = value.replace("–", "-").replace("—", "-")
    value = value.replace("[", "-").replace("]", "")

    tags: list[str] = []
    seen: set[str] = set()
    for chunk in re.split(r"[\n,;]+", value):
        for part in re.split(r"\s+-\s+|-", chunk):
            for tag in expand_tag_part(part):
                if not tag:
                    continue
                key = tag.casefold()
                if key not in seen:
                    seen.add(key)
                    tags.append(tag)
    return tags


def parse_docx(path: Path) -> list[TagRow]:
    document = Document(str(path))
    if not document.tables:
        raise RuntimeError("Tag file does not contain a table")

    rows: list[TagRow] = []
    for table_row in document.tables[0].rows[1:]:
        cells = [clean_text(cell.text) for cell in table_row.cells]
        if len(cells) < 3:
            continue
        tags = split_tag_cell(cells[2])
        if not tags:
            continue
        rows.append(
            TagRow(
                parasha=parasha_from_cell(cells[0]),
                title_text=cells[1],
                tags=tags,
            )
        )
    return rows


def article_score(row: TagRow, article: sqlite3.Row) -> float:
    row_title = normalize_for_match(row.title_text)
    title = normalize_for_match(article["title"])
    subtitle = normalize_for_match(article["subtitle"])
    combined = normalize_for_match(f"{article['title']} {article['subtitle'] or ''}")

    scores = [
        SequenceMatcher(None, row_title, combined).ratio(),
        SequenceMatcher(None, row_title, title).ratio(),
    ]
    if title and title in row_title:
        scores.append(0.99 if not subtitle or subtitle in row_title else 0.88)
    if subtitle and subtitle in row_title:
        scores.append(0.97)
    if row_title and row_title in combined:
        scores.append(0.96)

    score = max(scores)
    article_parasha = parasha_from_cell(article["parasha_name"])
    if row.parasha:
        if row.parasha == article_parasha:
            score += 0.08
        else:
            score -= 0.12
    return min(1.0, max(0.0, score))


def article_sort_key(row: TagRow, article: sqlite3.Row) -> tuple[float, int, float]:
    row_title = normalize_for_match(row.title_text)
    subtitle = normalize_for_match(article["subtitle"])
    combined = normalize_for_match(f"{article['title']} {article['subtitle'] or ''}")
    exact_subtitle = int(bool(subtitle and subtitle in row_title))
    combined_ratio = SequenceMatcher(None, row_title, combined).ratio()
    return (article_score(row, article), exact_subtitle, combined_ratio)


def sql_for_backend(sql: str, use_supabase: bool) -> str:
    if not use_supabase:
        return sql
    return sql.replace("?", "%s")


def execute(conn: Any, sql: str, params: tuple[Any, ...] = (), use_supabase: bool = False):
    return conn.execute(sql_for_backend(sql, use_supabase), params)


def first_value(row: Any) -> Any:
    if isinstance(row, dict):
        return next(iter(row.values()))
    return row[0]


def get_or_create_tag(conn: Any, name: str, use_supabase: bool) -> int:
    if use_supabase:
        execute(
            conn,
            "INSERT INTO public.tags (name, created_at) VALUES (?, ?) ON CONFLICT (name) DO NOTHING",
            (name, utc_now()),
            use_supabase,
        )
    else:
        execute(
            conn,
            "INSERT OR IGNORE INTO tags (name, created_at) VALUES (?, ?)",
            (name, utc_now()),
            use_supabase,
        )
    row = execute(conn, "SELECT id FROM tags WHERE name = ?", (name,), use_supabase).fetchone()
    if row is None:
        raise RuntimeError(f"Could not create tag: {name}")
    return int(row["id"])


def connect_database(use_supabase: bool) -> Any:
    if use_supabase:
        database_url = clean_env_value("DATABASE_URL")
        if not database_url:
            raise RuntimeError("DATABASE_URL is required when using --supabase")
        import psycopg
        from psycopg.rows import dict_row

        return psycopg.connect(database_url, row_factory=dict_row, prepare_threshold=None)

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def import_tags(
    rows: list[TagRow],
    clear_existing: bool = True,
    use_supabase: bool = False,
) -> tuple[int, int, list[str]]:
    with connect_database(use_supabase) as conn:
        articles = execute(
            conn,
            """
            SELECT articles.id, articles.title, articles.subtitle, articles.issue_number,
                   parashot.name AS parasha_name
            FROM articles
            JOIN parashot ON parashot.id = articles.parasha_id
            """,
            use_supabase=use_supabase,
        ).fetchall()

        if clear_existing:
            if use_supabase:
                execute(conn, "TRUNCATE public.article_tags, public.tags RESTART IDENTITY CASCADE", use_supabase=True)
            else:
                execute(conn, "DELETE FROM article_tags")
                execute(conn, "DELETE FROM tags")
                execute(conn, "DELETE FROM sqlite_sequence WHERE name IN ('tags')")

        matched_articles: set[int] = set()
        warnings: list[str] = []

        for row in rows:
            ranked = sorted(
                ((article_sort_key(row, article), article) for article in articles),
                key=lambda item: item[0],
                reverse=True,
            )
            if not ranked or ranked[0][0][0] < 0.78:
                warnings.append(f"No confident match: {row.parasha} | {row.title_text}")
                continue

            best_key, best_article = ranked[0]
            best_score = best_key[0]
            same_article_matches = [
                article
                for key, article in ranked
                if key[0] >= 0.9
                and article["parasha_name"] == best_article["parasha_name"]
                and article["title"] == best_article["title"]
                and (article["subtitle"] or "") == (best_article["subtitle"] or "")
            ]

            for article in same_article_matches:
                matched_articles.add(int(article["id"]))
                for tag in row.tags:
                    tag_id = get_or_create_tag(conn, tag, use_supabase)
                    execute(
                        conn,
                        (
                            "INSERT INTO public.article_tags (article_id, tag_id) VALUES (?, ?) ON CONFLICT DO NOTHING"
                            if use_supabase
                            else "INSERT OR IGNORE INTO article_tags (article_id, tag_id) VALUES (?, ?)"
                        ),
                        (article["id"], tag_id),
                        use_supabase,
                    )

            if best_score < 0.9:
                warnings.append(
                    f"Low confidence {best_score:.2f}: {row.parasha} | {row.title_text}"
                )

        linked_count = first_value(
            execute(conn, "SELECT COUNT(*) FROM article_tags", use_supabase=use_supabase).fetchone()
        )
        return len(matched_articles), int(linked_count), warnings


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--dry-run", action="store_true")
    parser.add_argument("--keep-existing", action="store_true")
    parser.add_argument("--supabase", action="store_true", help="Import directly into Supabase/Postgres using DATABASE_URL")
    args = parser.parse_args()

    if args.supabase and not clean_env_value("DATABASE_URL"):
        parser.error("DATABASE_URL is required when using --supabase")

    rows = parse_docx(args.docx)
    if args.dry_run:
        print(f"parsed_rows={len(rows)}")
        for row in rows[:10]:
            print(f"{row.parasha} | {row.title_text} | {', '.join(row.tags)}")
        return

    matched_articles, linked_count, warnings = import_tags(
        rows,
        clear_existing=not args.keep_existing,
        use_supabase=args.supabase,
    )
    print(f"parsed_rows={len(rows)} matched_articles={matched_articles} linked_tags={linked_count}")
    for warning in warnings[:20]:
        print(f"warning: {warning}")
    if len(warnings) > 20:
        print(f"warning: {len(warnings) - 20} more warnings")


if __name__ == "__main__":
    main()
